from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "research-narrative.md"
OUTPUT = ROOT / "research-narrative.docx"


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "DADCE0")
        borders.append(element)


def add_inline(paragraph, text):
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        run = paragraph.add_run(token[2:-2] if token.startswith("**") else token[1:-1])
        if token.startswith("**"):
            run.bold = True
        elif token.startswith("`"):
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        else:
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    for name, size, before, after, color in (
        ("Normal", 11, 0, 8, "000000"),
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[name].font.bold = False


def build():
    doc = Document()
    configure_styles(doc)
    lines = SOURCE.read_text().splitlines()
    i = 0
    first_heading = True
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[2:])
            run.font.name = "Arial"
            run.font.size = Pt(26)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = False
            first_heading = False
            i += 1
            continue
        if line.startswith("## "):
            add_inline(doc.add_paragraph(style="Heading 1"), line[3:])
            i += 1
            continue
        if line.startswith("### "):
            add_inline(doc.add_paragraph(style="Heading 2"), line[4:])
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[2:])
            run.italic = True
            i += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
            if match:
                caption, path = match.groups()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.add_run().add_picture(str(ROOT / path), width=Inches(6.35))
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(10)
                run = cp.add_run(caption)
                run.italic = True
                run.font.size = Pt(9.5)
            i += 1
            continue
        if line.startswith("$$"):
            math_lines = [line]
            if line.strip() == "$$":
                i += 1
                while i < len(lines):
                    math_lines.append(lines[i])
                    if lines[i].strip() == "$$":
                        break
                    i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            p.add_run("\n".join(math_lines))
            i += 1
            continue
        if line.startswith("| ") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            rows.pop(1)
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            set_table_borders(table)
            widths = [Inches(6.5 / len(rows[0]))] * len(rows[0])
            if len(rows[0]) == 2:
                widths = [Inches(4.9), Inches(1.6)]
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.width = widths[c_idx]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    add_inline(p, value)
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, line[2:])
            i += 1
            continue
        paragraph_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#|> |\!\[|\$\$|\| |-\s|\d+\.\s)", lines[i]
        ):
            paragraph_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
